import os
from docx import Document
from docx.shared import Pt, Inches
import textwrap

MAX_CHARACTERS = 35  # Maximum number of characters in a sentence
NUM_SPACES = 57  # Number of spaces in a blank line

def split_sentence(sentence):
    return textwrap.wrap(sentence, MAX_CHARACTERS)

def process_file(file_name):
    with open(file_name, "r", encoding="utf-8") as file:
        lines = file.readlines()
    
    processed_lines = []
    for line in lines:
        if line.strip() == '':
            processed_lines.append('')
        else:
            processed_lines.extend(split_sentence(line.strip()))

    return processed_lines

def set_font(paragraph, font_name, size):
    run = paragraph.runs[0]
    run.font.name = font_name
    run.font.size = Pt(size)

def set_spacing(paragraph, space_before, space_after):
    if paragraph._element.pPr is None:
        paragraph._element.get_or_add_pPr()
    p_spacing = paragraph._element.pPr.get_or_add_spacing()
    p_spacing.before = Pt(space_before)
    p_spacing.after = Pt(space_after)

def create_docx(sentences, docx_file_name):
    doc = Document()
    section = doc.sections[0]
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    counter = 0
    for sentence in sentences:
        if counter == 10:
            doc.add_page_break()
            counter = 0

        if sentence == '':
            paragraph = doc.add_paragraph()
        else:
            paragraph = doc.add_paragraph(sentence)
            set_font(paragraph, "KG Primary Penmanship Lined", 32)
            set_spacing(paragraph, 30, 0)
        
            blank_paragraph = doc.add_paragraph(' ' * NUM_SPACES)
            set_font(blank_paragraph, "KG Primary Penmanship Lined", 32)
            set_spacing(blank_paragraph, 30, 0)
            counter += 2

    doc.save(docx_file_name)

def main():
    directory = '.'  
    for file_name in os.listdir(directory):
        if file_name.endswith('.txt'):
            docx_file_name = os.path.splitext(file_name)[0] + '.docx'
            sentences = process_file(os.path.join(directory, file_name))
            create_docx(sentences, os.path.join(directory, docx_file_name))

if __name__ == '__main__':
    main()
